from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from app.services.document_parser import ParseError, parse_bytes, parse_file, parse_plain_text


def test_parse_plain_text_ok():
    text = "Alex Morgenstern\nMünchen\nErfahrung in Logistik und Operations seit vielen Jahren."
    r = parse_plain_text(text)
    assert "Alex Morgenstern" in r.text
    assert r.media_type == "txt"


def test_parse_plain_text_empty():
    with pytest.raises(ParseError) as ei:
        parse_plain_text("kurz")
    assert ei.value.code == "empty_or_scan"


def test_parse_txt_file(tmp_path: Path):
    p = tmp_path / "cv.txt"
    p.write_text(
        "Name: Alex Morgenstern\n"
        "Stadt: München\n"
        "Berufserfahrung als Leiter Logistik mit Verantwortung für Netz und OTIF.\n",
        encoding="utf-8",
    )
    r = parse_file(p)
    assert r.media_type == "txt"
    assert "Alex Morgenstern" in r.text


def test_parse_md_file(tmp_path: Path):
    p = tmp_path / "cv.md"
    p.write_text(
        "# Alex Morgenstern\n\n## Erfahrung\n\n- Geschäftsleitung Operations\n- OTIF und S&OP\n",
        encoding="utf-8",
    )
    r = parse_file(p)
    assert r.media_type == "md"
    assert "Geschäftsleitung" in r.text


def test_parse_docx_file(tmp_path: Path):
    p = tmp_path / "cv.docx"
    doc = Document()
    doc.add_heading("Alex Morgenstern", 0)
    doc.add_paragraph("München · alex@example.com")
    doc.add_paragraph("Geschäftsleitung Operations mit Verantwortung für OTIF, S&OP und 200 FTE.")
    doc.save(p)
    r = parse_file(p)
    assert r.media_type == "docx"
    assert "Alex Morgenstern" in r.text
    assert "OTIF" in r.text


def test_parse_unsupported():
    with pytest.raises(ParseError) as ei:
        parse_bytes(b"abc", "x.exe")
    assert ei.value.code == "unsupported_type"


def test_parse_empty_bytes():
    with pytest.raises(ParseError):
        parse_bytes(b"", "empty.txt")


def test_upload_txt_via_api(client, tmp_path: Path):
    content = (
        "Alex Morgenstern\nMünchen\n"
        "Leiter Logistik mit Erfahrung in Supply Chain und OTIF-Steuerung über Jahre.\n"
    ).encode()
    r = client.post(
        "/upload/cv",
        files={"file": ("master.txt", content, "text/plain")},
        headers={"Accept": "text/html"},
    )
    assert r.status_code == 200
    assert "Alex Morgenstern" in r.text
    assert "Text extrahiert" in r.text


def test_upload_empty_rejected(client):
    r = client.post(
        "/upload/cv",
        files={"file": ("empty.txt", b"", "text/plain")},
    )
    assert r.status_code in (413, 422)


def test_upload_paste(client):
    text = (
        "Alex Morgenstern aus München mit langjähriger Erfahrung in Operations, "
        "Logistiknetzwerken und Ergebnisverantwortung."
    )
    r = client.post(
        "/upload/cv",
        data={"paste_text": text},
        headers={"Accept": "text/html"},
    )
    assert r.status_code == 200
    assert "Alex Morgenstern" in r.text
