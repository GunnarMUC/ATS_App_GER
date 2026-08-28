from __future__ import annotations

from datetime import UTC, datetime
from pathlib import Path

from docx import Document
from fastapi import APIRouter, Depends, HTTPException, Request
from fastapi.responses import HTMLResponse, Response
from fastapi.templating import Jinja2Templates
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from sqlalchemy.orm import Session

from app.config import get_settings
from app.database import get_db
from app.models.orm import AdaptationPlan, FactLock, GeneratedDocument, JobDescription
from app.services import llm_client
from app.services.applications import link_document
from app.services.cover_generator import generate_cover
from app.services.cv_generator import generate_cv
from app.services.docx_builder import build_docx
from app.services.fact_guard import validate_cover_text, validate_generated_cv
from app.services.pdf_builder import build_pdf
from app.services.settings_service import ensure_settings_row
from app.services.text_builder import build_txt
from app.services.zip_export import build_application_zip

router = APIRouter(tags=["generate"])
templates = Jinja2Templates(directory="app/templates")


def _confirmed_plan(db: Session, job_id: str) -> AdaptationPlan:
    plan = (
        db.query(AdaptationPlan)
        .filter(AdaptationPlan.job_id == job_id, AdaptationPlan.status == "confirmed")
        .order_by(AdaptationPlan.confirmed_at.desc())
        .first()
    )
    if plan is None:
        raise HTTPException(status_code=409, detail="Kein bestätigter Plan.")
    return plan


@router.post("/jobs/{job_id}/generate")
async def generate_docs(
    job_id: str,
    request: Request,
    db: Session = Depends(get_db),
    type: str = "both",
):
    job = db.get(JobDescription, job_id)
    if job is None:
        raise HTTPException(status_code=404, detail="Stelle nicht gefunden.")
    if job.generating:
        raise HTTPException(status_code=409, detail="generating")
    plan = _confirmed_plan(db, job_id)
    lock = db.get(FactLock, plan.fact_lock_id)
    if lock is None or not lock.is_active:
        raise HTTPException(status_code=409, detail="FactLock nicht aktiv.")

    job.generating = True
    db.add(job)
    db.commit()
    created: list[str] = []
    try:
        if type in {"cv", "both"}:
            doc = await _generate_cv_doc(db, job, plan, lock)
            created.append(doc.id)
        if type in {"cover", "both"}:
            doc = await _generate_cover_doc(db, job, plan, lock)
            created.append(doc.id)
    finally:
        job.generating = False
        db.add(job)
        db.commit()

    if "text/html" in (request.headers.get("accept") or ""):
        return templates.TemplateResponse(
            request,
            "review.html",
            await _review_context(db, job, plan, lock),
        )
    return {"ok": True, "document_ids": created}


async def _generate_cv_doc(
    db: Session, job: JobDescription, plan: AdaptationPlan, lock: FactLock
) -> GeneratedDocument:
    cv_json = await generate_cv(
        lock.facts_json,
        plan.plan_json,
        job.analysis_json,
        db=db,
        use_llm=False,
    )
    guard = validate_generated_cv(lock.facts_json, cv_json)
    version = _next_version(db, job.id, "cv")
    settings = get_settings()
    base = settings.generated_dir / "cv" / f"{job.id}_v{version}"
    base.parent.mkdir(parents=True, exist_ok=True)
    docx_path = pdf_path = txt_path = None
    if guard.ok:
        docx_path = str(base.with_suffix(".docx"))
        pdf_path = str(base.with_suffix(".pdf"))
        txt_path = str(base.with_suffix(".txt"))
        build_docx(cv_json, Path(docx_path))
        build_pdf(cv_json, Path(pdf_path))
        Path(txt_path).write_text(build_txt(cv_json), encoding="utf-8")

    row = GeneratedDocument(
        type="cv",
        job_id=job.id,
        fact_lock_id=lock.id,
        plan_id=plan.id,
        version=version,
        structured_json={"cv": cv_json, "guard_errors": guard.errors},
        docx_path=docx_path,
        pdf_path=pdf_path,
        txt_path=txt_path,
        fact_guard_passed=guard.ok,
        created_at=datetime.now(UTC),
    )
    db.add(row)
    db.commit()
    db.refresh(row)
    link_document(db, job.id, row)
    db.refresh(row)
    return row


async def _generate_cover_doc(
    db: Session, job: JobDescription, plan: AdaptationPlan, lock: FactLock
) -> GeneratedDocument:
    text = await generate_cover(
        lock.facts_json,
        plan.plan_json,
        job.analysis_json,
        job.raw_text,
        db=db,
        use_llm=False,
    )
    guard = validate_cover_text(lock.facts_json, text)
    version = _next_version(db, job.id, "cover")
    settings = get_settings()
    base = settings.generated_dir / "cover" / f"{job.id}_v{version}"
    base.parent.mkdir(parents=True, exist_ok=True)
    docx_path = pdf_path = txt_path = None
    if guard.ok:
        txt_path = str(base.with_suffix(".txt"))
        Path(txt_path).write_text(text, encoding="utf-8")
        docx_path = str(base.with_suffix(".docx"))
        pdf_path = str(base.with_suffix(".pdf"))
        _write_cover_docx(text, Path(docx_path))
        _write_cover_pdf(text, Path(pdf_path))

    row = GeneratedDocument(
        type="cover",
        job_id=job.id,
        fact_lock_id=lock.id,
        plan_id=plan.id,
        version=version,
        structured_json={"text": text, "guard_errors": guard.errors},
        docx_path=docx_path,
        pdf_path=pdf_path,
        txt_path=txt_path,
        fact_guard_passed=guard.ok,
        created_at=datetime.now(UTC),
    )
    db.add(row)
    db.commit()
    db.refresh(row)
    link_document(db, job.id, row)
    db.refresh(row)
    return row


def _next_version(db: Session, job_id: str, doc_type: str) -> int:
    last = (
        db.query(GeneratedDocument)
        .filter(GeneratedDocument.job_id == job_id, GeneratedDocument.type == doc_type)
        .order_by(GeneratedDocument.version.desc())
        .first()
    )
    return (last.version + 1) if last else 1


def _write_cover_docx(text: str, path: Path) -> None:
    doc = Document()
    for line in text.splitlines():
        doc.add_paragraph(line)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))


def _write_cover_pdf(text: str, path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    c = canvas.Canvas(str(path), pagesize=A4)
    width, height = A4
    y = height - 50
    for line in text.splitlines():
        if y < 50:
            c.showPage()
            y = height - 50
        c.drawString(50, y, line[:110])
        y -= 14
    c.save()


async def _review_context(db, job, plan, lock):
    ensure_settings_row(db)
    health = await llm_client.check_health(db)
    cvs = (
        db.query(GeneratedDocument)
        .filter(GeneratedDocument.job_id == job.id, GeneratedDocument.type == "cv")
        .order_by(GeneratedDocument.version.desc())
        .all()
    )
    covers = (
        db.query(GeneratedDocument)
        .filter(GeneratedDocument.job_id == job.id, GeneratedDocument.type == "cover")
        .order_by(GeneratedDocument.version.desc())
        .all()
    )
    return {
        "health": health,
        "job": job,
        "plan": plan,
        "lock": lock,
        "cvs": cvs,
        "covers": covers,
        "latest_cv": cvs[0] if cvs else None,
        "latest_cover": covers[0] if covers else None,
    }


@router.get("/jobs/{job_id}/review", response_class=HTMLResponse)
async def review_page(job_id: str, request: Request, db: Session = Depends(get_db)):
    job = db.get(JobDescription, job_id)
    if job is None:
        raise HTTPException(status_code=404, detail="Stelle nicht gefunden.")
    plan = (
        db.query(AdaptationPlan)
        .filter(AdaptationPlan.job_id == job_id)
        .order_by(AdaptationPlan.id.desc())
        .first()
    )
    if plan is None:
        raise HTTPException(status_code=404, detail="Kein Plan.")
    lock = db.get(FactLock, plan.fact_lock_id)
    ctx = await _review_context(db, job, plan, lock)
    return templates.TemplateResponse(request, "review.html", ctx)


@router.get("/cv/compare/{job_id}", response_class=HTMLResponse)
async def compare_page(job_id: str, request: Request, db: Session = Depends(get_db)):
    job = db.get(JobDescription, job_id)
    if job is None:
        raise HTTPException(status_code=404, detail="Stelle nicht gefunden.")
    plan = (
        db.query(AdaptationPlan)
        .filter(AdaptationPlan.job_id == job_id)
        .order_by(AdaptationPlan.id.desc())
        .first()
    )
    if plan is None:
        raise HTTPException(status_code=404, detail="Kein Plan.")
    lock = db.get(FactLock, plan.fact_lock_id)
    if lock is None:
        raise HTTPException(status_code=404, detail="Kein FactLock.")
    ensure_settings_row(db)
    health = await llm_client.check_health(db)
    facts = lock.facts_json or {}
    plan_json = plan.plan_json or {}
    exp_map = {e["id"]: e for e in (facts.get("experience") or [])}
    skill_map = {s["id"]: s for s in (facts.get("skills") or [])}
    master_order = [e["id"] for e in (facts.get("experience") or [])]
    plan_order = list(plan_json.get("experience_order") or [])
    hidden = list(plan_json.get("hidden_experience_ids") or [])
    master_pos = {eid: i for i, eid in enumerate(master_order)}
    moved = [eid for i, eid in enumerate(plan_order) if master_pos.get(eid, i) != i]
    latest_cv = (
        db.query(GeneratedDocument)
        .filter(GeneratedDocument.job_id == job_id, GeneratedDocument.type == "cv")
        .order_by(GeneratedDocument.version.desc())
        .first()
    )
    guard_ok = bool(latest_cv and latest_cv.fact_guard_passed)
    guard_errors = []
    if latest_cv and latest_cv.structured_json:
        guard_errors = list((latest_cv.structured_json or {}).get("guard_errors") or [])
    return templates.TemplateResponse(
        request,
        "compare.html",
        {
            "health": health,
            "job": job,
            "plan": plan,
            "plan_json": plan_json,
            "facts": facts,
            "exp_map": exp_map,
            "skill_map": skill_map,
            "master_order": master_order,
            "plan_order": plan_order,
            "hidden": hidden,
            "moved": set(moved),
            "latest_cv": latest_cv,
            "guard_ok": guard_ok,
            "guard_errors": guard_errors,
        },
    )


@router.get("/documents/{doc_id}/download")
async def download_doc(doc_id: str, format: str = "docx", db: Session = Depends(get_db)):
    doc = db.get(GeneratedDocument, doc_id)
    if doc is None:
        raise HTTPException(status_code=404, detail="Dokument nicht gefunden.")
    if not doc.fact_guard_passed:
        raise HTTPException(status_code=409, detail="fact_guard_failed")
    fmt = format.lower()
    path = {"docx": doc.docx_path, "pdf": doc.pdf_path, "txt": doc.txt_path}.get(fmt)
    if not path or not Path(path).exists():
        raise HTTPException(status_code=404, detail="Datei fehlt.")
    data = Path(path).read_bytes()
    media = {
        "txt": "text/plain; charset=utf-8",
        "pdf": "application/pdf",
        "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    }[fmt]
    return Response(
        content=data,
        media_type=media,
        headers={"Content-Disposition": f'attachment; filename="{doc.type}_v{doc.version}.{fmt}"'},
    )


@router.get("/jobs/{job_id}/zip")
async def job_zip(job_id: str, db: Session = Depends(get_db)):
    job = db.get(JobDescription, job_id)
    if job is None:
        raise HTTPException(status_code=404, detail="Stelle nicht gefunden.")
    cv = (
        db.query(GeneratedDocument)
        .filter(
            GeneratedDocument.job_id == job_id,
            GeneratedDocument.type == "cv",
            GeneratedDocument.fact_guard_passed.is_(True),
        )
        .order_by(GeneratedDocument.version.desc())
        .first()
    )
    if cv is None:
        raise HTTPException(status_code=409, detail="Kein freigegebener CV.")
    cover = (
        db.query(GeneratedDocument)
        .filter(
            GeneratedDocument.job_id == job_id,
            GeneratedDocument.type == "cover",
            GeneratedDocument.fact_guard_passed.is_(True),
        )
        .order_by(GeneratedDocument.version.desc())
        .first()
    )
    plan = db.get(AdaptationPlan, cv.plan_id)
    lock = db.get(FactLock, cv.fact_lock_id)
    company = (job.analysis_json or {}).get("company") or "Firma"
    role = (plan.plan_json or {}).get("role_family") or "Rolle"
    z = build_application_zip(
        role=role,
        company=company,
        version=cv.version,
        cv_docx=Path(cv.docx_path).read_bytes(),
        cv_pdf=Path(cv.pdf_path).read_bytes(),
        cv_txt=Path(cv.txt_path).read_bytes(),
        cover_docx=Path(cover.docx_path).read_bytes() if cover and cover.docx_path else None,
        cover_pdf=Path(cover.pdf_path).read_bytes() if cover and cover.pdf_path else None,
        cover_txt=Path(cover.txt_path).read_text(encoding="utf-8")
        if cover and cover.txt_path
        else None,
        meta={
            "title": job.title_raw,
            "language": job.language,
            "hash": (lock.content_hash[:12] if lock else ""),
        },
    )
    return Response(
        content=z,
        media_type="application/zip",
        headers={"Content-Disposition": f'attachment; filename="Bewerbung_{job_id[:8]}.zip"'},
    )


@router.post("/documents/{doc_id}/edit")
async def edit_document(doc_id: str, request: Request, db: Session = Depends(get_db)):
    doc = db.get(GeneratedDocument, doc_id)
    if doc is None:
        raise HTTPException(status_code=404, detail="Dokument nicht gefunden.")
    body = await request.json()
    lock = db.get(FactLock, doc.fact_lock_id)
    if doc.type == "cover":
        text = body.get("text") or ""
        guard = validate_cover_text(lock.facts_json, text)
        data = dict(doc.structured_json or {})
        data["text"] = text
        data["guard_errors"] = guard.errors
        doc.structured_json = data
        doc.fact_guard_passed = guard.ok
        if guard.ok and doc.txt_path:
            Path(doc.txt_path).write_text(text, encoding="utf-8")
            if doc.docx_path:
                _write_cover_docx(text, Path(doc.docx_path))
            if doc.pdf_path:
                _write_cover_pdf(text, Path(doc.pdf_path))
    else:
        cv = body.get("cv") or body.get("facts")
        if not cv:
            raise HTTPException(status_code=422, detail="cv fehlt")
        guard = validate_generated_cv(lock.facts_json, cv)
        data = dict(doc.structured_json or {})
        data["cv"] = cv
        data["guard_errors"] = guard.errors
        doc.structured_json = data
        doc.fact_guard_passed = guard.ok
        if guard.ok:
            if doc.docx_path:
                build_docx(cv, Path(doc.docx_path))
            if doc.pdf_path:
                build_pdf(cv, Path(doc.pdf_path))
            if doc.txt_path:
                Path(doc.txt_path).write_text(build_txt(cv), encoding="utf-8")
    db.add(doc)
    db.commit()
    return {
        "ok": True,
        "fact_guard_passed": doc.fact_guard_passed,
        "errors": (doc.structured_json or {}).get("guard_errors"),
    }
