from __future__ import annotations

from io import BytesIO
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from app.services.cv_render import contact_line, format_month, load_cv, section_labels


def build_docx(facts: dict[str, Any], path: Path | None = None) -> bytes:
    cv = load_cv(facts)
    labels = section_labels(cv.language)
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    pf = style.paragraph_format
    pf.space_after = Pt(4)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE

    name_p = doc.add_paragraph()
    run = name_p.add_run(cv.personal.full_name)
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = "Arial"
    run.font.color.rgb = RGBColor(0x1A, 0x23, 0x32)

    contact = contact_line(cv)
    if contact:
        cp = doc.add_paragraph()
        cr = cp.add_run(contact)
        cr.font.size = Pt(10.5)
        cr.font.name = "Calibri"
        cr.font.color.rgb = RGBColor(0x1A, 0x23, 0x32)

    if cv.summary.strip():
        _heading(doc, labels["profile"])
        doc.add_paragraph(cv.summary.strip())

    _heading(doc, labels["experience"])
    for exp in cv.experience:
        head = ", ".join(x for x in (exp.title, exp.employer, exp.location) if x)
        hp = doc.add_paragraph()
        hr = hp.add_run(head)
        hr.bold = True
        hr.font.size = Pt(11)
        period = (
            f"{format_month(exp.start, cv.language)} – "
            f"{format_month(exp.end, cv.language)}"
        )
        pp = doc.add_paragraph()
        pr = pp.add_run(period)
        pr.font.size = Pt(10.5)
        pr.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        for b in exp.bullets:
            bp = doc.add_paragraph(b.text, style="List Bullet")
            for r in bp.runs:
                r.font.size = Pt(11)
                r.font.name = "Calibri"

    if cv.education:
        _heading(doc, labels["education"])
        for edu in cv.education:
            bits = [edu.degree, edu.institution, edu.field]
            doc.add_paragraph(", ".join(x for x in bits if x))
            if edu.start or edu.end:
                doc.add_paragraph(
                    f"{format_month(edu.start or '', cv.language)} – "
                    f"{format_month(edu.end or '', cv.language)}"
                )

    if cv.skills:
        _heading(doc, labels["skills"])
        doc.add_paragraph(", ".join(s.name for s in cv.skills))

    if cv.languages:
        _heading(doc, labels["languages"])
        doc.add_paragraph(", ".join(f"{lang.name} ({lang.level})" for lang in cv.languages))

    if cv.certifications:
        _heading(doc, labels["certifications"])
        for c in cv.certifications:
            bit = c.name
            if c.year:
                bit += f" ({c.year})"
            if c.issuer:
                bit += f" — {c.issuer}"
            doc.add_paragraph(bit, style="List Bullet")

    # ensure no tables
    assert not doc.tables

    buf = BytesIO()
    doc.save(buf)
    data = buf.getvalue()
    if path is not None:
        path.parent.mkdir(parents=True, exist_ok=True)
        path.write_bytes(data)
    return data


def _heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(12)
    r.font.name = "Arial"
    r.font.color.rgb = RGBColor(0x1A, 0x23, 0x32)
