from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path


ALLOWED_EXTENSIONS = {".pdf", ".docx", ".md", ".txt"}
MIN_TEXT_CHARS = 20


class ParseError(Exception):
    def __init__(self, message: str, code: str = "parse_error") -> None:
        super().__init__(message)
        self.message = message
        self.code = code


@dataclass
class ParseResult:
    text: str
    media_type: str
    warnings: list[str]


def extension_of(filename: str) -> str:
    return Path(filename).suffix.lower()


def media_type_for(ext: str) -> str:
    return {
        ".pdf": "pdf",
        ".docx": "docx",
        ".md": "md",
        ".txt": "txt",
    }.get(ext, "txt")


def parse_file(path: Path, original_filename: str | None = None) -> ParseResult:
    name = original_filename or path.name
    ext = extension_of(name)
    if ext not in ALLOWED_EXTENSIONS:
        raise ParseError(
            f"Dateityp nicht erlaubt: {ext or '(ohne)'}. Erlaubt: pdf, docx, md, txt.",
            code="unsupported_type",
        )
    if not path.exists():
        raise ParseError("Datei nicht gefunden.", code="not_found")

    if ext == ".pdf":
        text, warnings = _parse_pdf(path)
    elif ext == ".docx":
        text, warnings = _parse_docx(path)
    else:
        text, warnings = _parse_text(path)

    cleaned = _normalize_text(text)
    if len(cleaned.strip()) < MIN_TEXT_CHARS:
        raise ParseError(
            "Aus dieser Datei ließ sich kein brauchbarer Text lesen. "
            "Gescannte PDFs werden in v1 nicht unterstützt. Bitte DOCX oder Text.",
            code="empty_or_scan",
        )
    return ParseResult(text=cleaned, media_type=media_type_for(ext), warnings=warnings)


def parse_bytes(data: bytes, filename: str) -> ParseResult:
    import tempfile

    ext = extension_of(filename)
    if ext not in ALLOWED_EXTENSIONS:
        raise ParseError(
            f"Dateityp nicht erlaubt: {ext or '(ohne)'}. Erlaubt: pdf, docx, md, txt.",
            code="unsupported_type",
        )
    with tempfile.NamedTemporaryFile(suffix=ext, delete=False) as tmp:
        tmp.write(data)
        tmp_path = Path(tmp.name)
    try:
        return parse_file(tmp_path, original_filename=filename)
    finally:
        tmp_path.unlink(missing_ok=True)


def parse_plain_text(text: str, media_type: str = "txt") -> ParseResult:
    cleaned = _normalize_text(text)
    if len(cleaned.strip()) < MIN_TEXT_CHARS:
        raise ParseError(
            "Text zu kurz oder leer.",
            code="empty_or_scan",
        )
    return ParseResult(text=cleaned, media_type=media_type, warnings=[])


def _normalize_text(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    lines = [ln.rstrip() for ln in text.split("\n")]
    return "\n".join(lines).strip()


def _parse_text(path: Path) -> tuple[str, list[str]]:
    raw = path.read_bytes()
    for enc in ("utf-8", "utf-8-sig", "latin-1", "cp1252"):
        try:
            return raw.decode(enc), []
        except UnicodeDecodeError:
            continue
    return raw.decode("utf-8", errors="replace"), ["encoding_fallback"]


def _parse_docx(path: Path) -> tuple[str, list[str]]:
    from docx import Document

    doc = Document(str(path))
    parts: list[str] = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            parts.append(t)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    warnings = ["docx_has_tables"] if doc.tables else []
    return "\n".join(parts), warnings


def _parse_pdf(path: Path) -> tuple[str, list[str]]:
    warnings: list[str] = []
    text = ""
    try:
        import fitz

        with fitz.open(str(path)) as doc:
            chunks: list[str] = []
            for page in doc:
                chunks.append(page.get_text("text") or "")
            text = "\n".join(chunks)
    except Exception:
        warnings.append("pymupdf_failed")
        text = ""

    if len(text.strip()) < MIN_TEXT_CHARS:
        try:
            import pdfplumber

            with pdfplumber.open(str(path)) as pdf:
                chunks = []
                for page in pdf.pages:
                    chunks.append(page.extract_text() or "")
                text = "\n".join(chunks)
                warnings.append("pdfplumber_fallback")
        except Exception:
            warnings.append("pdfplumber_failed")

    return text, warnings
