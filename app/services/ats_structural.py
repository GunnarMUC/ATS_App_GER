from __future__ import annotations

from pathlib import Path
from typing import Any

from app.models.schemas import CVStructure


def analyze_source_file(path: Path | None, media_type: str, raw_text: str = "") -> dict[str, Any]:
    issues: list[dict[str, str]] = []
    notes: list[str] = []

    if path is not None and path.exists():
        if media_type == "docx":
            issues.extend(_scan_docx(path))
        elif media_type == "pdf":
            issues.extend(_scan_pdf(path))
        elif media_type in {"txt", "md"}:
            notes.append("Reintext — gut für ATS-Paste, Layout-Risiken entfallen.")
    else:
        notes.append("Originaldatei nicht geprüft (nur gespeicherte Fakten).")

    if media_type == "pdf" and raw_text is not None and len(raw_text.strip()) < 40:
        issues.append(
            {
                "code": "text_in_image",
                "severity": "block",
                "message_de": "Sehr wenig Text — möglicherweise Scan oder Bild-PDF.",
            }
        )

    score = _score(issues)
    return {
        "schema_version": "1.0",
        "score_hint": score,
        "issues": issues,
        "notes_de": notes,
    }


def analyze_cv_structure(facts: dict[str, Any] | CVStructure) -> dict[str, Any]:
    if isinstance(facts, CVStructure):
        cv = facts
    else:
        cv = CVStructure.model_validate(facts)

    issues: list[dict[str, str]] = []
    notes: list[str] = [
        "Report auf Basis der gesperrten/gespeicherten Fakten (Ausgabe-Builder, einspaltig).",
    ]

    p = cv.personal
    contact_bits = [p.email, p.phone, p.city]
    if not any(contact_bits):
        issues.append(
            {
                "code": "missing_contact",
                "severity": "warn",
                "message_de": "Wenig Kontaktangaben (E-Mail/Telefon/Stadt).",
            }
        )
    else:
        notes.append("Kontaktfelder sind im Dokumentkörper vorgesehen (nicht nur Kopfzeile).")

    if p.birth_date or p.marital_status or p.nationality:
        notes.append(
            "Klassische Personalien liegen in den Fakten; Default-Output lässt sie weg (AGG/ATS)."
        )

    score = _score(issues)
    if not issues:
        score = max(score, 90)

    return {
        "schema_version": "1.0",
        "score_hint": score,
        "issues": issues,
        "notes_de": notes,
    }


def _score(issues: list[dict[str, str]]) -> int:
    score = 100
    for issue in issues:
        sev = issue.get("severity")
        if sev == "block":
            score -= 35
        elif sev == "warn":
            score -= 15
        else:
            score -= 5
    return max(0, min(100, score))


def _scan_docx(path: Path) -> list[dict[str, str]]:
    issues: list[dict[str, str]] = []
    try:
        from docx import Document
        from docx.oxml.ns import qn

        doc = Document(str(path))
        if doc.tables:
            issues.append(
                {
                    "code": "table_layout",
                    "severity": "warn",
                    "message_de": f"DOCX enthält {len(doc.tables)} Tabelle(n) — ATS kann Zellenreihenfolge falsch lesen.",
                }
            )
        # images in document body
        blips = doc.element.body.findall(".//" + qn("a:blip"))
        if blips:
            issues.append(
                {
                    "code": "image_or_photo",
                    "severity": "warn",
                    "message_de": "Bilder/Grafiken im DOCX erkannt.",
                }
            )
        # text boxes / shapes (w:txbxContent)
        txbx = doc.element.body.findall(".//" + qn("w:txbxContent"))
        if txbx:
            issues.append(
                {
                    "code": "textboxes",
                    "severity": "warn",
                    "message_de": "Textboxen erkannt — für ATS ungeeignet.",
                }
            )
        # headers with content
        for section in doc.sections:
            header_text = "\n".join(p.text for p in section.header.paragraphs).strip()
            if header_text and any(
                x in header_text.lower() for x in ("@", "tel", "http", "linkedin", "+")
            ):
                issues.append(
                    {
                        "code": "header_contact_only",
                        "severity": "info",
                        "message_de": "Kontakt scheint in der Kopfzeile zu stehen — besser im Body.",
                    }
                )
                break
    except Exception:
        issues.append(
            {
                "code": "table_layout",
                "severity": "info",
                "message_de": "DOCX-Strukturprüfung teilweise fehlgeschlagen.",
            }
        )
    return issues


def _scan_pdf(path: Path) -> list[dict[str, str]]:
    issues: list[dict[str, str]] = []
    try:
        import fitz

        with fitz.open(str(path)) as doc:
            has_image = False
            tiny = False
            multi = False
            for page in doc:
                if page.get_images():
                    has_image = True
                blocks = page.get_text("dict").get("blocks") or []
                x_centers: list[float] = []
                for b in blocks:
                    if b.get("type") != 0:
                        continue
                    for line in b.get("lines") or []:
                        for span in line.get("spans") or []:
                            size = float(span.get("size") or 0)
                            if 0 < size < 8:
                                tiny = True
                            bbox = span.get("bbox") or b.get("bbox")
                            if bbox:
                                x_centers.append((bbox[0] + bbox[2]) / 2)
                if x_centers:
                    left = sum(1 for x in x_centers if x < page.rect.width * 0.42)
                    right = sum(1 for x in x_centers if x > page.rect.width * 0.58)
                    if left > 5 and right > 5:
                        multi = True
            if has_image:
                issues.append(
                    {
                        "code": "image_or_photo",
                        "severity": "warn",
                        "message_de": "PDF enthält Bilder.",
                    }
                )
            if tiny:
                issues.append(
                    {
                        "code": "tiny_font",
                        "severity": "warn",
                        "message_de": "Sehr kleine Schrift erkannt (< 8 pt).",
                    }
                )
            if multi:
                issues.append(
                    {
                        "code": "multi_column",
                        "severity": "warn",
                        "message_de": "Vermutlich mehrspaltiges Layout.",
                    }
                )
    except Exception:
        issues.append(
            {
                "code": "text_in_image",
                "severity": "info",
                "message_de": "PDF-Strukturprüfung teilweise fehlgeschlagen.",
            }
        )
    return issues
